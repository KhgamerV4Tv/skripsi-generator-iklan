import io
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_document_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

def set_document_footer(doc):
    for section in doc.sections:
        footer = section.footer
        
        # Clear existing paragraphs if any
        for paragraph in footer.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)

        # Create a new paragraph for the footer
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # We can simulate the page number in center by using tabs, but for simplicity
        # and reliable standard format, we can just align right and put the text.
        # Actually, MS Word requires field codes for real page numbers.
        # Let's add the static text first.
        run = p.add_run("Univ Kristen Petra")
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True

def set_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

def create_docx(text_content, image_bytes_list=None):
    doc = Document()
    
    # Apply standard styles
    set_font(doc)
    set_document_margins(doc)
    set_document_footer(doc)
    
    # Add content
    for paragraph in text_content.split('\n'):
        p = doc.add_paragraph(paragraph)
        # Apply font specifically to run if style doesn't cascade perfectly
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            
    # Add images if provided
    if image_bytes_list:
        for img_bytes in image_bytes_list:
            image_stream = io.BytesIO(img_bytes)
            try:
                # Assuming images are somewhat standard sized, adding them directly
                doc.add_picture(image_stream, width=Cm(14))
                doc.add_paragraph()
            except Exception as e:
                pass # Skip image if corrupted or unsupported format

    # Save to BytesIO stream
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def test_format():
    print("Testing DOCX formatter...")
    output = create_docx("This is a test paragraph for UK Petra standard format.")
    assert output is not None
    assert output.getvalue().startswith(b"PK")
    print("Format test passed!")
    return True
